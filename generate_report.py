"""Generate the project report as a Word document."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

RESULTS_DIR = os.path.join(os.path.dirname(__file__), "results")
REPORT_PATH = os.path.join(os.path.dirname(__file__), "report", "Project3_Group2_Report.docx")
os.makedirs(os.path.dirname(REPORT_PATH), exist_ok=True)

doc = Document()

# ---- Style setup ----
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.space_before = Pt(10)
    hs.paragraph_format.space_after = Pt(4)
    if level == 1:
        hs.font.size = Pt(14)
    elif level == 2:
        hs.font.size = Pt(12)
    else:
        hs.font.size = Pt(11)

# ---- Margins ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_para(text, bold=False, italic=False, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_figure(image_path, caption, width=Inches(4.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.italic = True
    cap.paragraph_format.space_after = Pt(8)


# ===================================================================
# TITLE
# ===================================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(
    "Sequential Sentence Classification in Medical Abstracts\n"
    "Using LSTM and BioBERT"
)
title_run.font.name = "Times New Roman"
title_run.font.size = Pt(16)
title_run.bold = True
title_p.paragraph_format.space_after = Pt(6)

# Authors
author_p = doc.add_paragraph()
author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author_p.add_run(
    "Jialu Liang, Benjamin Tondre, Qing Wang\n"
)
author_run.font.name = "Times New Roman"
author_run.font.size = Pt(11)
author_p.paragraph_format.space_after = Pt(2)

course_p = doc.add_paragraph()
course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
course_run = course_p.add_run(
    "BME 6938 — Spring 2026 — Project 3 — Group 2"
)
course_run.font.name = "Times New Roman"
course_run.font.size = Pt(11)
course_p.paragraph_format.space_after = Pt(2)

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info_p.add_run(
    "GitHub: https://github.com/wq2581/bme6938-2026-spring-project3-group2"
)
info_run.font.name = "Times New Roman"
info_run.font.size = Pt(10)
info_p.paragraph_format.space_after = Pt(12)


# ===================================================================
# ABSTRACT
# ===================================================================
doc.add_heading("Abstract", level=1)
add_para(
    "Automated classification of sentences in medical abstracts into rhetorical roles "
    "(Background, Objective, Methods, Results, Conclusions) can facilitate systematic reviews, "
    "literature screening, and evidence-based medicine. In this work, we address this task using "
    "the PubMed RCT 20k dataset, which contains approximately 180,000 training sentences drawn "
    "from randomized controlled trial abstracts. We implement and compare two neural approaches: "
    "(1) a bidirectional LSTM baseline trained from scratch with learned word embeddings, and "
    "(2) a fine-tuned BioBERT transformer model pre-trained on biomedical corpora. "
    "On the held-out test set, the LSTM baseline achieves 82.6% accuracy and 76.1% macro F1, "
    "while BioBERT achieves 86.9% accuracy and 80.5% macro F1, representing a 4.3 percentage-point "
    "improvement in accuracy. Both models perform well on majority classes (Methods, Results) "
    "but struggle with minority classes (Objective, Background), suggesting that class imbalance "
    "and inherent linguistic ambiguity remain challenges. We discuss clinical implications, "
    "limitations, and directions for future work."
)

# ===================================================================
# 1. INTRODUCTION
# ===================================================================
doc.add_heading("1. Introduction", level=1)
add_para(
    "The exponential growth of biomedical literature poses a significant challenge for clinicians "
    "and researchers who must efficiently identify relevant evidence from published studies. "
    "PubMed alone indexes over 35 million citations, with thousands of new articles added daily [1]. "
    "Randomized controlled trials (RCTs) are the gold standard for evaluating clinical interventions, "
    "and their abstracts follow a structured format organized into rhetorical sections: Background, "
    "Objective, Methods, Results, and Conclusions [2]. However, many abstracts are published in "
    "unstructured free-text form, making it difficult to quickly extract specific information such "
    "as study findings or methodology."
)
add_para(
    "Automatic sequential sentence classification (SSC) in medical abstracts addresses this problem "
    "by assigning each sentence to its rhetorical role [3]. This capability has direct clinical "
    "applications: it can accelerate systematic review workflows by enabling automated PICO "
    "(Population, Intervention, Comparison, Outcome) extraction [4], support clinical decision "
    "support systems by surfacing relevant evidence by section type, and assist medical students "
    "in understanding the standard organization of scientific writing [5]."
)
add_para(
    "In this project, we tackle the sentence classification task using the PubMed RCT 20k dataset [3] "
    "and compare two neural architectures. First, we implement a bidirectional LSTM baseline that "
    "learns word representations from scratch, serving as a straightforward sequence model. Second, "
    "we fine-tune BioBERT [6], a transformer model pre-trained on PubMed abstracts and PMC full-text "
    "articles, which captures domain-specific biomedical knowledge. Our goal is to evaluate how "
    "transfer learning from biomedical corpora compares to training a simpler recurrent model "
    "from scratch on this clinically relevant task."
)

# ===================================================================
# 2. LITERATURE REVIEW
# ===================================================================
doc.add_heading("2. Related Works", level=1)
add_para(
    "Sequential sentence classification in medical abstracts has been studied extensively. "
    "Dernoncourt and Lee [3] introduced the PubMed RCT dataset and established baselines using "
    "convolutional neural networks (CNNs) and recurrent neural networks (RNNs), achieving around "
    "90% accuracy with models that leveraged both token and character embeddings along with "
    "sentence position features. Their work demonstrated that deep learning methods significantly "
    "outperform traditional feature-engineered approaches such as SVMs and CRFs for this task."
)
add_para(
    "The introduction of transformer-based language models fundamentally changed NLP. "
    "Devlin et al. [7] proposed BERT, a bidirectional transformer pre-trained on masked language "
    "modeling and next-sentence prediction, which achieved state-of-the-art results across diverse "
    "NLP benchmarks. Recognizing the domain gap between general and biomedical text, Lee et al. [6] "
    "developed BioBERT by continuing BERT's pre-training on PubMed abstracts (4.5B words) and PMC "
    "full-text articles (13.5B words). BioBERT demonstrated significant improvements over BERT on "
    "biomedical NER, relation extraction, and question answering tasks."
)
add_para(
    "Subsequent domain-specific models have further advanced biomedical NLP. Gu et al. [8] "
    "introduced PubMedBERT, pre-trained from scratch exclusively on PubMed text, arguing that "
    "domain-specific pre-training from scratch outperforms mixed-domain continual pre-training. "
    "Alsentzer et al. [9] developed BioClinicalBERT by further pre-training BioBERT on clinical "
    "notes from MIMIC-III, targeting clinical NLP applications."
)
add_para(
    "For the specific task of abstract sentence classification, Jin and Szolovits [10] explored "
    "hierarchical models that capture both sentence-level and document-level context, showing that "
    "inter-sentence dependencies improve classification accuracy. Cohan et al. [11] proposed "
    "pretrained models for scientific documents (SciBERT) trained on semantic scholar papers. "
    "Beltagy et al. [12] introduced Longformer for handling long documents efficiently."
)
add_para(
    "Hochreiter and Schmidhuber [13] proposed Long Short-Term Memory (LSTM) networks to address "
    "the vanishing gradient problem in standard RNNs, enabling effective modeling of long-range "
    "dependencies in sequential data. Bidirectional LSTMs [14] extend this by processing sequences "
    "in both forward and backward directions, capturing contextual information from both sides of "
    "each token. Vaswani et al. [15] introduced the Transformer architecture based on self-attention, "
    "which has since become the foundation for modern NLP models."
)
add_para(
    "Our work builds on these foundations by directly comparing an LSTM baseline against a fine-tuned "
    "BioBERT model on the PubMed RCT 20k dataset, providing a clear illustration of the performance "
    "gap between traditional recurrent models and pre-trained transformers in biomedical text "
    "classification."
)

# ===================================================================
# 3. METHODS & DATA
# ===================================================================
doc.add_heading("3. Methods and Data", level=1)

doc.add_heading("3.1 Dataset", level=2)
add_para(
    "We use the PubMed RCT 20k dataset [3], available on Hugging Face (armanc/pubmed-rct20k). "
    "The dataset contains sentences extracted from approximately 20,000 RCT abstracts from PubMed, "
    "each labeled with one of five rhetorical roles: Background, Objective, Methods, Results, and "
    "Conclusions. The dataset provides pre-defined train/validation/test splits: approximately "
    "176,642 training sentences, 29,672 validation sentences, and 29,578 test sentences."
)
add_para(
    "The class distribution is imbalanced. Methods (33.4%) and Results (32.8%) dominate, followed "
    "by Conclusions (15.4%), Background (10.4%), and Objective (7.9%). This imbalance reflects the "
    "natural structure of RCT abstracts, where methodology and findings typically constitute the "
    "longest sections."
)

doc.add_heading("3.2 Text Preprocessing", level=2)
add_para(
    "For the LSTM model, we apply text preprocessing including lowercasing, replacing digits with "
    "a NUM token, removing punctuation, and collapsing whitespace. We build a vocabulary from the "
    "training set with a minimum frequency threshold of 2, yielding approximately 36,000 unique "
    "tokens. Unknown words are mapped to a special UNK token, and sequences are padded or truncated "
    "to a maximum length of 128 tokens."
)
add_para(
    "For BioBERT, we use the model's native WordPiece tokenizer (dmis-lab/biobert-base-cased-v1.2) "
    "with padding to max_length=128 and truncation enabled. No additional text normalization is "
    "applied, as the pre-trained tokenizer handles subword segmentation of biomedical terminology."
)

doc.add_heading("3.3 Model Architectures", level=2)
add_para(
    "LSTM Baseline. We implement a bidirectional LSTM classifier with the following architecture: "
    "an embedding layer (128 dimensions), followed by a 2-layer bidirectional LSTM (256 hidden units "
    "per direction), dropout (p=0.3), and a linear classification head mapping the concatenated "
    "final forward and backward hidden states (512 dimensions) to 5 output classes. The model has "
    "approximately 5.4 million trainable parameters."
)
add_para(
    "BioBERT Transformer. We fine-tune BioBERT (dmis-lab/biobert-base-cased-v1.2) [6], which "
    "consists of 12 transformer layers, 768 hidden dimensions, and 12 attention heads "
    "(approximately 110 million parameters). A linear classification head is added on top of the "
    "[CLS] token representation. The pre-trained weights capture biomedical language patterns from "
    "PubMed and PMC corpora."
)

doc.add_heading("3.4 Training Details", level=2)
add_para(
    "LSTM training. We train for 10 epochs using Adam optimizer with learning rate 1e-3, batch "
    "size 128, and gradient clipping at max_norm=1.0. A ReduceLROnPlateau scheduler with patience 2 "
    "and factor 0.5 is used. The best model is selected based on validation macro F1 score."
)
add_para(
    "BioBERT training. We fine-tune for 5 epochs using AdamW optimizer with learning rate 2e-5, "
    "weight decay 0.01, warmup ratio 0.1, and batch size 32. Mixed-precision (FP16) training is "
    "enabled on GPU. Early stopping with patience 2 based on validation F1 is applied. The best "
    "checkpoint is selected based on validation macro F1."
)
add_para(
    "Both models are trained on the same data splits to ensure a fair comparison. All experiments "
    "are conducted on a single NVIDIA L4 GPU (16 GB)."
)

# ===================================================================
# 4. RESULTS & EVALUATION
# ===================================================================
doc.add_heading("4. Results and Evaluation", level=1)

doc.add_heading("4.1 Overall Performance", level=2)
add_para(
    "Table 1 summarizes the test set performance of both models. BioBERT outperforms the LSTM "
    "baseline across all metrics, with a 4.3 percentage-point improvement in accuracy and a 4.4 "
    "percentage-point improvement in macro F1."
)

# Table 1: Overall performance
cap_t1 = doc.add_paragraph()
cap_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_t1 = cap_t1.add_run("Table 1: Overall test set performance comparison.")
r_t1.font.name = "Times New Roman"
r_t1.font.size = Pt(10)
r_t1.bold = True

table1 = doc.add_table(rows=3, cols=5)
table1.style = "Light Grid Accent 1"
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ["Model", "Accuracy", "Macro Precision", "Macro Recall", "Macro F1"]
row0 = table1.rows[0].cells
for i, h in enumerate(headers):
    row0[i].text = h
row1 = table1.rows[1].cells
for i, v in enumerate(["BiLSTM", "0.826", "0.777", "0.754", "0.761"]):
    row1[i].text = v
row2 = table1.rows[2].cells
for i, v in enumerate(["BioBERT", "0.869", "0.821", "0.803", "0.805"]):
    row2[i].text = v
for row in table1.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
doc.add_paragraph()

doc.add_heading("4.2 Per-Class Performance", level=2)
add_para(
    "Table 2 shows per-class F1 scores. Both models perform best on Methods and Results, which "
    "are the most frequent classes. The largest performance gap between the two models appears in "
    "the minority classes: BioBERT improves F1 on Background by 4.0 points, on Objective by 1.8 "
    "points, and on Conclusions by 2.2 points compared to the LSTM baseline."
)

# Table 2: Per-class F1
cap_t2 = doc.add_paragraph()
cap_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_t2 = cap_t2.add_run("Table 2: Per-class F1 scores on the test set.")
r_t2.font.name = "Times New Roman"
r_t2.font.size = Pt(10)
r_t2.bold = True

table2 = doc.add_table(rows=3, cols=6)
table2.style = "Light Grid Accent 1"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = ["Model", "Background", "Objective", "Methods", "Results", "Conclusions"]
for i, h in enumerate(h2):
    table2.rows[0].cells[i].text = h
for i, v in enumerate(["BiLSTM", "0.638", "0.630", "0.906", "0.883", "0.748"]):
    table2.rows[1].cells[i].text = v
for i, v in enumerate(["BioBERT", "0.678", "0.648", "0.926", "0.900", "0.770"]):
    table2.rows[2].cells[i].text = v
for row in table2.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
doc.add_paragraph()

doc.add_heading("4.3 Confusion Matrices", level=2)
add_para(
    "Figures 1 and 2 show confusion matrices for both models. The most common confusion pattern "
    "is between Background and Objective, which is expected since both categories describe the "
    "study context and motivation. Another notable confusion is between Results and Conclusions, "
    "as concluding sentences often restate key findings."
)

add_figure(
    os.path.join(RESULTS_DIR, "rnn_confusion_matrix.png"),
    "Figure 1: Confusion matrix for the BiLSTM model on the test set.",
    width=Inches(4.2),
)
add_figure(
    os.path.join(RESULTS_DIR, "transformer_confusion_matrix.png"),
    "Figure 2: Confusion matrix for the BioBERT model on the test set.",
    width=Inches(4.2),
)

doc.add_heading("4.4 Training Dynamics", level=2)
add_para(
    "Figures 3 and 4 show training curves. The LSTM model exhibits clear overfitting: training "
    "loss continues to decrease while validation loss increases after epoch 2, and validation "
    "accuracy plateaus around 83%. The BioBERT model also shows some overfitting, with validation "
    "loss increasing after epoch 1 while training loss decreases steadily. Early stopping selects "
    "the best checkpoint before overfitting degrades generalization."
)

add_figure(
    os.path.join(RESULTS_DIR, "rnn_training_curves.png"),
    "Figure 3: Training and validation loss/accuracy curves for the BiLSTM model.",
    width=Inches(5.0),
)
add_figure(
    os.path.join(RESULTS_DIR, "transformer_training_curves.png"),
    "Figure 4: Training and validation loss curves for the BioBERT model.",
    width=Inches(5.0),
)

# ===================================================================
# 5. DISCUSSION & LIMITATIONS
# ===================================================================
doc.add_heading("5. Discussion and Limitations", level=1)

doc.add_heading("5.1 Interpretation of Results", level=2)
add_para(
    "The BioBERT model consistently outperforms the LSTM baseline across all classes, confirming "
    "the value of transfer learning from biomedical corpora for this task. The pre-trained "
    "representations capture domain-specific terminology and contextual patterns that the LSTM "
    "must learn entirely from the training data. Notably, the performance gap is most pronounced "
    "for underrepresented classes (Background, Objective), where the limited training examples "
    "make it harder for the LSTM to learn robust representations."
)
add_para(
    "The confusion between Background and Objective is clinically understandable: both describe "
    "the rationale and goals of a study, and the boundary between establishing context and stating "
    "an aim can be linguistically subtle. Similarly, Conclusions sentences that restate quantitative "
    "findings are easily confused with Results, reflecting genuine ambiguity in rhetorical structure."
)

doc.add_heading("5.2 Clinical Relevance", level=2)
add_para(
    "Automated sentence classification in medical abstracts has practical applications in clinical "
    "workflows. A model that reliably identifies Results and Conclusions sentences could help "
    "clinicians rapidly extract actionable evidence from the literature during time-constrained "
    "clinical decision-making. Integration with systematic review tools could reduce the manual "
    "effort required in evidence synthesis, which currently represents a significant bottleneck "
    "in evidence-based medicine [4]."
)
add_para(
    "Both models achieve over 88% F1 on Methods and Results, the two most clinically relevant "
    "categories for understanding what was done and what was found. However, the lower performance "
    "on Background and Objective (F1 around 0.63-0.68) suggests that current models are less "
    "reliable for identifying contextual and motivational content, which may limit their utility "
    "in fully automated abstract structuring systems."
)

doc.add_heading("5.3 Limitations", level=2)
add_para(
    "Dataset limitations. The PubMed RCT 20k dataset only covers RCT abstracts, which tend to "
    "follow a relatively standardized structure. Generalization to other study designs (observational "
    "studies, case reports, meta-analyses) or to full-text articles remains untested. The class "
    "imbalance in the dataset, with Methods and Results comprising over 66% of all sentences, may "
    "bias models toward majority classes."
)
add_para(
    "Model limitations. Our models classify sentences independently without considering their "
    "position in the abstract or relationships with neighboring sentences. Prior work has shown "
    "that sequential context (e.g., a Background sentence is more likely to precede an Objective "
    "sentence) provides valuable signal [3, 10]. Incorporating such hierarchical or sequential "
    "modeling could improve performance, particularly for distinguishing between related categories."
)
add_para(
    "Overfitting. Both models show signs of overfitting, particularly the LSTM model whose "
    "validation loss increases substantially after early epochs. More aggressive regularization, "
    "data augmentation, or reduced model capacity could mitigate this issue."
)

doc.add_heading("5.4 Ethical Considerations", level=2)
add_para(
    "Deploying NLP models in clinical settings requires careful consideration of potential biases. "
    "The training data is drawn exclusively from English-language PubMed abstracts, which may "
    "underrepresent research from non-English-speaking regions and certain medical specialties. "
    "Misclassification of sentence roles could lead to incorrect evidence extraction, which in a "
    "clinical decision support context could potentially influence patient care. Any deployment "
    "should include human oversight and validation against domain-expert annotations."
)

doc.add_heading("5.5 Future Work", level=2)
add_para(
    "Several directions could improve upon this work: (1) incorporating sentence position and "
    "inter-sentence context through hierarchical models or CRF layers; (2) experimenting with "
    "other domain-specific models such as PubMedBERT [8] or SciBERT [11]; (3) addressing class "
    "imbalance through oversampling, class-weighted loss, or focal loss; (4) evaluating on "
    "additional datasets to test cross-domain generalization; and (5) integrating attention "
    "visualization for model interpretability in clinical applications."
)

# ===================================================================
# AUTHOR CONTRIBUTIONS
# ===================================================================
doc.add_heading("Author Contributions", level=1)
add_para(
    "Jialu Liang: Data preprocessing pipeline, RNN/LSTM baseline model implementation, "
    "EDA notebook, training curve analysis."
)
add_para(
    "Qing Wang: BioBERT transformer fine-tuning, evaluation framework, demo notebook, "
    "error analysis and visualization."
)
add_para(
    "Benjamin Tondre: Literature review, report writing, results documentation."
)

# ===================================================================
# REFERENCES
# ===================================================================
doc.add_heading("References", level=1)

refs = [
    "[1] National Library of Medicine. PubMed overview. https://pubmed.ncbi.nlm.nih.gov, 2024.",
    "[2] ICMJE. Recommendations for the Conduct, Reporting, Editing, and Publication of Scholarly Work in Medical Journals. ICMJE, 2023.",
    "[3] F. Dernoncourt and J. Y. Lee. PubMed 200k RCT: a dataset for sequential sentence classification in medical abstracts. In Proc. IJCNLP, 2017.",
    "[4] A. M. Cohen et al. Reducing workload in systematic review preparation using automated citation classification. JAMIA, 13(2):206-219, 2006.",
    "[5] B. Nye et al. A corpus with multi-level annotations of patients, interventions and outcomes to support language processing for medical literature. In Proc. ACL, 2018.",
    "[6] J. Lee et al. BioBERT: a pre-trained biomedical language representation model for biomedical text mining. Bioinformatics, 36(4):1234-1240, 2020.",
    "[7] J. Devlin et al. BERT: Pre-training of deep bidirectional transformers for language understanding. In Proc. NAACL-HLT, 2019.",
    "[8] Y. Gu et al. Domain-specific language model pretraining for biomedical natural language processing. ACM Trans. Comput. Healthcare, 3(1):1-23, 2022.",
    "[9] E. Alsentzer et al. Publicly available clinical BERT embeddings. In Proc. ClinicalNLP Workshop, NAACL, 2019.",
    "[10] D. Jin and P. Szolovits. Hierarchical neural networks for sequential sentence classification in medical scientific abstracts. In Proc. EMNLP, 2018.",
    "[11] A. Cohan et al. Pretrained language models for sequential sentence classification. In Proc. EMNLP, 2019.",
    "[12] I. Beltagy et al. Longformer: The long-document transformer. arXiv:2004.05150, 2020.",
    "[13] S. Hochreiter and J. Schmidhuber. Long short-term memory. Neural Computation, 9(8):1735-1780, 1997.",
    "[14] M. Schuster and K. K. Paliwal. Bidirectional recurrent neural networks. IEEE Trans. Signal Processing, 45(11):2673-2681, 1997.",
    "[15] A. Vaswani et al. Attention is all you need. In Proc. NeurIPS, 2017.",
    "[16] T. Wolf et al. Transformers: State-of-the-art natural language processing. In Proc. EMNLP: System Demonstrations, 2020.",
    "[17] A. Paszke et al. PyTorch: An imperative style, high-performance deep learning library. In Proc. NeurIPS, 2019.",
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

# ---- Save ----
doc.save(REPORT_PATH)
print(f"Report saved to {REPORT_PATH}")
